import os
import re #regex
    
from collections import OrderedDict
from docx2python import docx2python
from io import StringIO
from docx import Document
from docx.shared import Inches


from django.conf import settings
from django.db.models import Count
from django.db.models import F
from django.db.models.functions import Length
from django.core.management.base import BaseCommand

from ...models import Section, Neighborhood, Category, Blob


# virtualenv venvironment
# cd..
# pip install -r requirements.txt

# c:\Python\bookbuild-main\bookbuild_src\
# venvironment\Scripts\activate
# python manage.py export_text 

class Command(BaseCommand):


    def handle(self, *args, **options):
        kwargs = {}

        filename = "export.docx"
        filepath = os.path.join(settings.BASE_DIR, "importer", "export", filename)
        

        document = Document()
        
        #walk through neighborhoods sorted by blob count
 #       for n in Neighborhood.objects.annotate(blob_count=Count('blob')).filter(blob_count__gte=1).order_by('-blob_count'):
 
        #walk through neighborhoods sorted by order
        for n in Neighborhood.objects.order_by('order'):
            
            
            
           # if n.title in ["Silom", "Verkehr", "Festivals", "Apps", "Sukhumvit", "Chinatown & Altstadt", "Innenstadt", "Metropolregion", "Ausflug", "Chatuchak & Ratchada", "Upper Sukhumvit", "Siam & Pratunam"] : continue #only 1 chapter for now 
            
 #           if not (n.title == "Reiseinfos"): continue
            
            
            
            
            #print out neighborhood heading
            document.add_heading(n.title, 0)
            
            
        
            
            
#            #sort sections by sum of priority      
#            section = OrderedDict({"None":[]})
#            for b in Blob.objects.filter(neighborhood=n, priority__isnull=False).order_by('section__title'):
#                if not b.category.section:
#                    section["None"] = [b]
#                elif b.category.section.title not in section:
#                    section[b.category.section.title] = [b]
#                else:
#                    section[b.category.section.title].append(b)



         
#            section_accumulated_priority = {}
#            for c, b_list in section.items():
#                accumulated_priority = sum([x.priority for x in b_list if x.priority is not None])
#                if accumulated_priority in section_accumulated_priority:
#                    section_accumulated_priority[accumulated_priority].append(c)
#                else:
#                    section_accumulated_priority[accumulated_priority] = [c]
#            sorted_priorities = list(section_accumulated_priority.keys())
#            sorted_priorities.sort(reverse=True)

#           for p in sorted_priorities:
#                sections = section_accumulated_priority.get(p)
#                for c in sections:
             
            #Neighborhood Description first
            for b in Blob.objects.filter(neighborhood=n):
                if (b.category != None and b.category.title == "Kapitelbeschreibung"): document.add_paragraph(b.main_text)
            document.add_page_break()
            
            
            #Excursions get a special treatment
            if (n.title=="Ausflug"):
            
                for c in Section.objects.order_by('order'):
                    if (c.title in ["Halbtagesausflug","Ganztagesausflug","Mehrtagesausflug",]): #and Blob.objects.filter(neighborhood=n, priority__gte=2, priority__lte=4).filter(section=c).order_by('-priority').count()>0:
                    
                        document.add_heading(c.title, level=1)
                        
                        #walk through blobs in section Ausflug
                        for b in Blob.objects.filter(neighborhood=n, priority__gte=2, section=c).order_by('-priority'):
                            
                            document.add_heading(b.title, level=2)
                            document.add_paragraph(f"{str(b.category_text)}")
    #                        document.add_paragraph(f"Priority: {str(b.priority)}")
                            document.add_paragraph(b.main_text)
                            document.add_paragraph(stripHTML(b.footer_text))
     #                       document.add_paragraph("----")
            
                continue #excursions done
            
                    
                      
            
            
            #walk through Prio 5 items next
            document.add_heading("Highlights", level=1)
            for b in Blob.objects.filter(neighborhood=n, priority__gte=5).order_by(F('order').asc(nulls_last=True)):
            
                document.add_heading(b.title, level=2)
                document.add_paragraph(f"{str(b.category_text)}")
#                document.add_paragraph(f"Priority: {str(b.priority)}")
                document.add_paragraph(b.main_text)
                document.add_paragraph(stripHTML(b.footer_text))  
                document.add_page_break()
            
            
            
            
            #walk through sections sorted by section.order
            for c in Section.objects.order_by('order'):
                
                
            #    if (c.title=="Routen" or c.title=="Fernverkehr"): continue
              
                                    
                if Blob.objects.filter(neighborhood=n, priority__gte=2, priority__lte=4).filter(category__section=c).count()>0: 
                    
                    
    #                        if (b.category != None and b.category.section.title == "Routen"): document.add_page_break()
                    
                    document.add_heading(c.title, level=1)
                    #walk through blobs in section
                    for b in Blob.objects.filter(neighborhood=n, priority__gte=2, priority__lte=4).filter(category__section=c).order_by('-priority'): #.order_by(F('order').asc(nulls_last=True)):
                
                         # sort blobs by priority
        #                b_list = section.get(c)
        #                if len(b_list) > 0:
        #                    b_list.sort(key=lambda x: x.priority, reverse=True)
        #                    for b in b_list:
                        document.add_heading(b.title, level=2)
                                #title.add_run('bold').bold = True
        #                            document.add_paragraph(f"{str(b.category.section)}")
                        document.add_paragraph(f"{str(b.category_text)}")
    #                        document.add_paragraph(f"Priority: {str(b.priority)}")
                        document.add_paragraph(b.main_text)
                        document.add_paragraph(stripHTML(b.footer_text))
    #                   document.add_paragraph("----")

            document.add_page_break()
            
        
        
        
        
        
            #repeat for transit hubs 
            # for s in Section.objects.filter(title="Fernverkehr"):
                # document.add_heading(s.title, level=1) 
                
                # #walk through categories
# #                for c in Category.objects.filter(section=s):
                
                # #walk through blobs in section Fernverkehr
                # for b in Blob.objects.filter(category__section=s, priority__gte=2).order_by('-priority', -Length('main_text')):
                    
                    # document.add_heading(b.title, level=2)
                    # document.add_paragraph(f"{str(b.category_text)}")
# #                        document.add_paragraph(f"Priority: {str(b.priority)}")
                    # document.add_paragraph(b.main_text)
                    # document.add_paragraph(stripHTML(b.footer_text))
# #                       document.add_paragraph("----")
        
            

        document.save(filepath)
        
        
            
    
        #repeat for booklet
        filename_booklet = "export_booklet.docx"
        filepath_booklet = os.path.join(settings.BASE_DIR, "importer", "export", filename_booklet)
      
        document_booklet = Document()
        
        for n in Neighborhood.objects.order_by('order'):
        
            for c in Section.objects.order_by('order'):
                if (c.title!="Routen"): continue
                for b in Blob.objects.filter(neighborhood=n, priority__gte=2).filter(category__section=c).order_by('-priority'):
 #               for b in Blob.objects.filter(neighborhood=n).filter(category__section__title="Routen"):
                        
                    document_booklet.add_heading(b.title, level=2)
                    document_booklet.add_paragraph(f"{str(b.category_text)}")
        #          document.add_paragraph(f"Priority: {str(b.priority)}")
                    document_booklet.add_paragraph(b.main_text)
                    document_booklet.add_paragraph(stripHTML(b.footer_text))  
                    document_booklet.add_page_break()
                        
            
        document_booklet.save(filepath_booklet)
        
        
        
        
        #repeat headers only Prio 2-5
        filename_header = "export_header.docx"
        filepath_header = os.path.join(settings.BASE_DIR, "importer", "export", filename_header)
      
        document_header = Document()
        
        for n in Neighborhood.objects.order_by('order'):
            document_header.add_heading(n.title, 0)
            for c in Section.objects.order_by('order'):
                if Blob.objects.filter(neighborhood=n, priority__gte=2, priority__lte=5).filter(category__section=c).count()>0: 
                    document_header.add_heading(c.title, level=1)
                
                    for b in Blob.objects.filter(neighborhood=n, priority__gte=2, priority__lte=5).filter(category__section=c).order_by('-priority'):
                        document_header.add_heading(b.title, level=2)
                        
            
            document_header.add_page_break()
        document_header.save(filepath_header)
        
        
        


def stripHTML(text):
    if (text is None) : return text
    return re.sub('<[^<]+?>', '', text)
