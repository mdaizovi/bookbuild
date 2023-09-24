from docx import Document
from docx2python import docx2python
from io import StringIO
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .models import Section, Neighborhood, Category, Blob


def import_text(filename):
    doc_result = docx2python(f"{filename}.docx")
    # get separate components of the document
    # print(doc_result.body)

    # get the text from Zen of Python
    # print(len(doc_result.body))
    # print(len(doc_result.body[0][0]))
    body = doc_result.body[0][0]
    print(len(body[-1]))
    # print(body[-1])

    # with open(f"{filename}.docx", 'rb') as f:
    #     doc = Document(f)
    #     for d in doc.paragraphs:

    #         print(f"-{d.text}-")
    # print(type(d)) # <class 'docx.text.paragraph.Paragraph'>
    # print(d.runs) # a run is basically a list of items that are on same line.

    # for r in d.runs:
    #     print(f"-{r.text}-")

    # lastd = doc.paragraphs[-5]
    # #print(lastd.text)
    # for r in lastd.runs:
    #     print(f"{r.text}")

    # source_stream = StringIO(f.read())
    #     source_stream = StringIO(f.read().decode('latin-1'))

    # document = Document(source_stream)
    # print("document")
    # print(document)
    # print("\n\n-----------\n\n")
    # print(document.splitlines())
    # source_stream.close()


filename = "Texte-2017/03-ausflug"
import_text(filename)
